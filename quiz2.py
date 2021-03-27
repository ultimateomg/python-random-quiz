from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches, Pt
import random
from docx.enum.text import WD_LINE_SPACING
import argparse

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Create a ArcHydro schema')
    parser.add_argument('--filename', metavar='path', required=True,
                        help='File de thi')
    parser.add_argument('--number', metavar='path', required=True,
                        help='So luong de thi')
    args = parser.parse_args()
    lst_idx_questions = [0]
    lst_idx_answers = {0: []}
    doc = Document(args.filename)

    last_idx = 0
    for i in range(len(doc.paragraphs)):
        t = doc.paragraphs[i].text
        if (t.strip() == ""):
            last_idx = i+1
            lst_idx_answers[last_idx] = []
            lst_idx_questions.append(last_idx)
            continue
        if i not in lst_idx_questions:
            lst_idx_answers[last_idx].append(i)
    print(lst_idx_questions, lst_idx_answers)

    total_anwser = []
    _max = 0  # Số lượng câu trả lời nhiều nhất
    list_correct_answers = {}
    for k in lst_idx_answers.keys():
        _max = max(len(lst_idx_answers[k]), _max)
        for j in range(len(lst_idx_answers[k])):
            if j == 0:
                list_correct_answers[k] = lst_idx_answers[k][0]
            total_anwser.append(j)
    print(_max, total_anwser, list_correct_answers)

    num_quizzes = int(args.number)
    list_quizzes = []
    ans_idx = ["a) ", "b) ", "c) ", "d) ", "e) ", "f) "]
    labels = ["A", "B", "C", "D", "E", "F", "G", "H"]

    doc_answered = Document()
    table = doc_answered.add_table(
        rows=len(lst_idx_questions)+1, cols=num_quizzes+1)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = "Câu"
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True

    for q in range(len(lst_idx_questions)):
        table.rows[q+1].cells[0].text = str(q+1)
        table.rows[q+1].cells[0].paragraphs[0].runs[0].font.bold = True

    for idx in range(num_quizzes):
        table.rows[0].cells[idx+1].text = "Đề "+str(idx+1)
        table.rows[0].cells[idx+1].paragraphs[0].runs[0].font.bold = True

        random.shuffle(lst_idx_questions)
        while ' '.join(str(s) for s in lst_idx_questions) in list_quizzes:
            random.shuffle(lst_idx_questions)
        list_quizzes.append(' '.join(str(s) for s in lst_idx_questions))
        _counter = [0 for _ in range(_max)]

        while True:
            _counter = [0 for _ in range(_max)]
            _correct = []
            for q in lst_idx_questions:
                random.shuffle(lst_idx_answers[q])
                _counter[lst_idx_answers[q].index(
                    list_correct_answers[q])] += 1
                _correct.append(
                    labels[lst_idx_answers[q].index(list_correct_answers[q])])
            __max = max(_counter)/_max
            __min = min(_counter)/_max
            if __max-__min < (2.1/_max):
                break
        print("Bo de", idx, ": ", lst_idx_questions, _counter, lst_idx_answers)
        quiz = Document()
        style = quiz.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        for _idx, question in enumerate(lst_idx_questions):
            # print(doc.paragraphs[question].text)
            _p = quiz.add_paragraph("", style=style)
            _p.paragraph_format.space_before = Pt(0)
            _p.paragraph_format.space_after = Pt(0)
            _p.paragraph_format.line_spacing = 1
            run = _p.add_run("Câu "+str(_idx+1)+": ")
            run.bold = True
            for run in doc.paragraphs[question].runs:
                _r = _p.add_run(run.text)
                _r.bold = run.bold
                _r.italic = run.italic
                _r.underline = run.underline
                _r.font.color.rgb = run.font.color.rgb
                _r.style.name = run.style.name
            _p.paragraph_format.alignment = _p.paragraph_format.alignment
            for _idx2, answer in enumerate(lst_idx_answers[question]):
                _p = quiz.add_paragraph(ans_idx[_idx2], style=style)
                _p.paragraph_format.space_before = Pt(0)
                _p.paragraph_format.space_after = Pt(0)
                _p.paragraph_format.line_spacing = 1
                for run in doc.paragraphs[answer].runs:
                    _r = _p.add_run(run.text)
                    _r.bold = run.bold
                    _r.italic = run.italic
                    _r.underline = run.underline
                    _r.font.color.rgb = run.font.color.rgb
                    _r.style.name = run.style.name
                _p.paragraph_format.alignment = _p.paragraph_format.alignment
        for _idx3, _c in enumerate(_correct):
            row = table.rows[_idx3+1]
            row.cells[idx+1].text = _c
        print("answered: ", _correct)
        quiz.save(str(idx)+".docx")

    doc_answered.save("Dap-an.docx")
