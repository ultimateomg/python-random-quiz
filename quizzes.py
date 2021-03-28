from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.text.run import Run
from docx.oxml.text.run import CT_R
from docx.shared import Inches, Pt
import random
from docx.enum.text import WD_LINE_SPACING
import argparse
import copy
from io import BytesIO
from docx.enum.shape import WD_INLINE_SHAPE
import shutil
import zipfile
import os
import fnmatch


def zip_directory(folder_path, zip_path):
    with zipfile.ZipFile(zip_path, mode='w') as zipf:
        len_dir_path = len(folder_path)
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, file_path[len_dir_path:])


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Create a ArcHydro schema')
    parser.add_argument('--filename', metavar='path', required=True,
                        help='File de thi')
    parser.add_argument('--number', metavar='path', required=True,
                        help='So luong de thi')
    args = parser.parse_args()
    lst_idx_questions = [0]
    lst_idx_answers = {0: []}
    doc = Document(args.filename)
    out_dir1 = "temporary/tmp1"
    with zipfile.ZipFile(args.filename, 'r') as zip_ref:
        zip_ref.extractall(out_dir1)
    media_path_1 = os.path.join(out_dir1, "word", "media")
    _rels_1 = os.path.join(out_dir1, "word", "_rels")
    last_idx = 0
    for i in range(len(doc.paragraphs)):
        t = doc.paragraphs[i].text
        if (t.strip() == ""):
            last_idx = i+1
            lst_idx_answers[last_idx] = []
            lst_idx_questions.append(last_idx)
            continue
        if i not in lst_idx_questions:
            lst_idx_answers[last_idx].append(i)
    print(lst_idx_questions, lst_idx_answers)

    total_anwser = []
    _max = 0  # Số lượng câu trả lời nhiều nhất
    list_correct_answers = {}
    for k in lst_idx_answers.keys():
        _max = max(len(lst_idx_answers[k]), _max)
        for j in range(len(lst_idx_answers[k])):
            if j == 0:
                list_correct_answers[k] = lst_idx_answers[k][0]
            total_anwser.append(j)
    print(_max, total_anwser, list_correct_answers)

    num_quizzes = int(args.number)
    list_quizzes = []
    ans_idx = ["a) ", "b) ", "c) ", "d) ", "e) ", "f) "]
    labels = ["A", "B", "C", "D", "E", "F", "G", "H"]

    doc_answered = Document()
    table = doc_answered.add_table(
        rows=len(lst_idx_questions)+1, cols=num_quizzes+1)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = "Câu"
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True

    for q in range(len(lst_idx_questions)):
        table.rows[q+1].cells[0].text = str(q+1)
        table.rows[q+1].cells[0].paragraphs[0].runs[0].font.bold = True

    for idx in range(num_quizzes):
        table.rows[0].cells[idx+1].text = "Đề "+str(idx+1)
        table.rows[0].cells[idx+1].paragraphs[0].runs[0].font.bold = True

        random.shuffle(lst_idx_questions)
        while ' '.join(str(s) for s in lst_idx_questions) in list_quizzes:
            random.shuffle(lst_idx_questions)
        list_quizzes.append(' '.join(str(s) for s in lst_idx_questions))
        _counter = [0 for _ in range(_max)]

        while True:
            _counter = [0 for _ in range(_max)]
            _correct = []
            for q in lst_idx_questions:
                random.shuffle(lst_idx_answers[q])
                _counter[lst_idx_answers[q].index(
                    list_correct_answers[q])] += 1
                _correct.append(
                    labels[lst_idx_answers[q].index(list_correct_answers[q])])
            __max = max(_counter)/_max
            __min = min(_counter)/_max
            if __max-__min < (2.1/_max):
                break
        print("Bo de", idx, ": ", lst_idx_questions, _counter, lst_idx_answers)
        quiz = Document()
        style = quiz.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        for _idx, question in enumerate(lst_idx_questions):
            doc2 = copy.deepcopy(doc)
            quiz.element.body.append(doc2.element.body[question])
            if len(quiz.paragraphs[-1].runs) > 0:
                new_run_element = quiz.paragraphs[-1]._element._new_r()
                _current_run = quiz.paragraphs[-1].runs[0]
                _current_run._element.addprevious(
                    new_run_element)
                new_run = Run(new_run_element, _current_run._parent)
                new_run.text = "Câu " + str(_idx+1)+": "
                new_run.bold = True
            for _idx2, answer in enumerate(lst_idx_answers[question]):
                doc3 = copy.deepcopy(doc)
                quiz.element.body.append(doc3.element.body[answer])
                if len(quiz.paragraphs[-1].runs) > 0:
                    new_run_element = quiz.paragraphs[-1]._element._new_r()
                    _current_run = quiz.paragraphs[-1].runs[0]
                    _current_run._element.addprevious(
                        new_run_element)
                    new_run = Run(new_run_element, _current_run._parent)
                    new_run.text = ans_idx[_idx2]
            del doc2, doc3
        for _idx3, _c in enumerate(_correct):
            row = table.rows[_idx3+1]
            row.cells[idx+1].text = _c
        print("answered: ", _correct)
        style = quiz.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        for _p in quiz.paragraphs:
            _p.paragraph_format.space_before = Pt(0)
            _p.paragraph_format.space_after = Pt(0)
            _p.paragraph_format.line_spacing = 1
            _p.style = style
            for _r in _p.runs:
                _font = _r.font
                _font.name = 'Times New Roman'
                _font.size = Pt(13)
        quiz.save(str(idx+1)+".docx")
        out_dir2 = "temporary/tmp2"
        with zipfile.ZipFile(str(idx+1)+".docx", 'r') as zip_ref:
            zip_ref.extractall(out_dir2)

        media_path_2 = os.path.join(out_dir2, "word", "media")
        _rels_2 = os.path.join(out_dir2, "word", "_rels")
        if os.path.exists(_rels_2):
            shutil.rmtree(_rels_2)
        if os.path.exists(media_path_2):
            shutil.rmtree(media_path_2)
        shutil.copytree(media_path_1, media_path_2)
        shutil.copytree(_rels_1, _rels_2)
        zip_directory(out_dir2, str(idx+1)+".docx")
    doc_answered.save("Dap-an.docx")
